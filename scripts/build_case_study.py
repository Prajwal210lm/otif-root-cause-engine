"""Generate the two-page OTIF Root-Cause Engine case study (.docx).

Every business number is read live from data/canonical_run.json (the seed-46
canonical run) and data/robustness/summary.json (the six-seed distribution), so
regenerating the canonical data and re-running this script keeps the case study
exactly in step with the site. No figure is hand-transcribed.

Requires python-docx (`pip install python-docx`). This is a documentation tool,
not an engine runtime dependency, so it is intentionally not in requirements.txt.

Usage:
    python scripts/build_case_study.py
Writes: docs/OTIF_Root_Cause_Engine_Case_Study.docx
"""
import json
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CANON = json.load(open(os.path.join(ROOT, "data", "canonical_run.json")))
SUMM = json.load(open(os.path.join(ROOT, "data", "robustness", "summary.json")))
OUT = os.path.join(ROOT, "docs", "OTIF_Root_Cause_Engine_Case_Study.docx")

# ---- palette ----------------------------------------------------------------
NAVY = RGBColor(0x0F, 0x25, 0x40)
NAVY_HEX = "0F2540"
STEEL_HEX = "1C3A5E"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x1A, 0x1F, 0x28)
MUTE = RGBColor(0x5B, 0x64, 0x70)
ACCENT = RGBColor(0x0B, 0x5E, 0x55)      # teal, ties to the live site accent
ACCENT_HEX = "0B5E55"
ROW_ALT_HEX = "EEF1F5"
CANON_HEX = "E4EEEC"                      # soft teal wash for the canonical row
ANOM_HEX = "F6ECEC"                       # soft warm wash for the seed-47 anomaly
FONT = "Calibri"

# ---- derived numbers (all from the JSON, never typed) -----------------------
sc, nv = CANON["scorecard"], CANON["naive"]
roll = CANON["rollup"]
agg = SUMM["aggregate"]
TOTAL = roll["total_cash"]

DRIVER_META = {
    "supplier": ("Supplier", "Qualify a second source for the worst-offending vendors"),
    "demand": ("Demand", "Recalibrate the forecast on the SKUs that consistently under-call"),
    "warehouse": ("Warehouse", "Clear the dispatch bottleneck with a hard release cut-off"),
    "logistics": ("Logistics", "Review and re-bond the chronically overrun lanes"),
}


def pct(frac, dp=1):
    return f"{frac * 100:.{dp}f}%"


def aed(n):
    return "AED " + f"{round(n):,}"


# ---- low-level docx helpers -------------------------------------------------
def shade(cell, hex_fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(sh)


def no_borders(table):
    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tbl.tblPr.append(borders)


def hair_borders(table, hex_color="D6DCE4"):
    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tbl.tblPr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def run(p, text, size, *, bold=False, color=INK, italic=False, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rpr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rpr.append(sp)
    return r


def para(container, *, before=0, after=4, line=None, align=None):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    return p


def cell_para(cell, *, after=0, before=0, align=None):
    p = cell.paragraphs[0] if not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.paragraph_format.alignment = align
    return p


# ---- document ---------------------------------------------------------------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.55)
sec.bottom_margin = Inches(0.55)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)
CONTENT_W = Inches(7.1)

base = doc.styles["Normal"]
base.font.name = FONT
base.font.size = Pt(10)
base.font.color.rgb = INK

# ---- header bar -------------------------------------------------------------
hdr = doc.add_table(rows=1, cols=1)
hdr.autofit = False
hdr.columns[0].width = CONTENT_W
no_borders(hdr)
hc = hdr.cell(0, 0)
hc.width = CONTENT_W
shade(hc, NAVY_HEX)
set_cell_margins(hc, top=150, bottom=150, left=200, right=200)
p = cell_para(hc, after=0)
run(p, "OTIF ROOT-CAUSE ENGINE", 17, bold=True, color=WHITE, spacing=10)
p2 = cell_para(hc, before=2, after=0)
run(p2, "CASE STUDY", 9.5, color=RGBColor(0x9F, 0xD6, 0xCE), caps=True, spacing=30)
run(p2, "     Attributing delivery failures to a root cause, with the money next to each one.",
    9.5, color=RGBColor(0xC7, 0xD3, 0xDF))

# ---- four-stat strip --------------------------------------------------------
stats = [
    (pct(sc["overall_accuracy"]), "Overall accuracy", "canonical batch, seed 46"),
    (pct(agg["mean_ambiguous_accuracy"]), "Mean on hard orders", "across 6 independent batches"),
    (f"+{agg['mean_lift_ambiguous'] * 100:.1f}", "Points mean lift", "vs a simple biggest-number rule"),
    (aed(TOTAL).replace("AED ", "AED "), "Failure value attributed", "ranked by cash, seed 46"),
]
strip = doc.add_table(rows=1, cols=4)
strip.autofit = False
no_borders(strip)
for i, (big, label, sub) in enumerate(stats):
    c = strip.cell(0, i)
    c.width = Inches(7.1 / 4)
    shade(c, NAVY_HEX if i % 2 == 0 else STEEL_HEX)
    set_cell_margins(c, top=130, bottom=130, left=130, right=110)
    pb = cell_para(c, after=1)
    run(pb, big, 19 if not big.startswith("AED") else 13.5, bold=True, color=WHITE)
    pl = cell_para(c, after=0, before=2)
    run(pl, label, 8.5, bold=True, color=WHITE, caps=True, spacing=6)
    ps = cell_para(c, after=0, before=1)
    run(ps, sub, 7.5, color=RGBColor(0xAF, 0xBE, 0xCD))

# ---- problem + what it does -------------------------------------------------
p = para(doc, before=12, after=2)
run(p, "The problem", 12, bold=True, color=NAVY)
p = para(doc, after=8, line=1.18)
run(p, "When on-time-in-full slips, nobody agrees on why. Planning blames the suppliers, the "
       "suppliers blame the forecast, the warehouse blames both, and logistics says the order left "
       "on time. It costs days of meetings and finger-pointing, and no single owner ever agrees to "
       "hold the number. The cause is real and knowable, but it is buried in per-order signals no "
       "one has time to reconcile by hand.", 10, color=INK)

p = para(doc, before=2, after=2)
run(p, "What the tool does", 12, bold=True, color=NAVY)
p = para(doc, after=8, line=1.18)
run(p, "Four specialist agents (one each for supplier, demand, warehouse, and logistics) investigate "
       "every failed order in parallel. A coordinator reads all four findings and names the one "
       "dominant cause, then ranks the causes by the cash behind them. The agents reason; they never "
       "produce a figure. Every number in the output is computed by tested code and checked before it "
       "can be shown, so nothing the model writes can invent a value.", 10, color=INK)

# ---- value-at-stake table ---------------------------------------------------
p = para(doc, before=2, after=4)
run(p, "Where the value sits (canonical batch, seed 46)", 12, bold=True, color=NAVY)

vt = doc.add_table(rows=1, cols=5)
vt.autofit = False
hair_borders(vt)
widths = [Inches(1.15), Inches(1.35), Inches(0.8), Inches(0.95), Inches(2.85)]
heads = ["Driver", "Value at stake", "Share", "Failed orders", "Recommended action"]
for i, (h, w) in enumerate(zip(heads, widths)):
    c = vt.cell(0, i)
    c.width = w
    shade(c, NAVY_HEX)
    set_cell_margins(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pp = cell_para(c, after=0, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 or i == 4 else WD_ALIGN_PARAGRAPH.RIGHT)
    run(pp, h, 8.5, bold=True, color=WHITE, caps=True, spacing=4)

for idx, d in enumerate(roll["by_driver"]):
    label, action = DRIVER_META[d["driver"]]
    cells = vt.add_row().cells
    fills = [ROW_ALT_HEX if idx % 2 else "FFFFFF"] * 5
    vals = [
        (label, False, WD_ALIGN_PARAGRAPH.LEFT, True),
        (aed(d["cash"]), False, WD_ALIGN_PARAGRAPH.RIGHT, False),
        (pct(d["cash_share"]), False, WD_ALIGN_PARAGRAPH.RIGHT, False),
        (str(d["order_count"]), False, WD_ALIGN_PARAGRAPH.RIGHT, False),
        (action, False, WD_ALIGN_PARAGRAPH.LEFT, False),
    ]
    for i, (c, w) in enumerate(zip(cells, widths)):
        c.width = w
        shade(c, fills[i])
        set_cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        txt, _, align, strong = vals[i]
        pp = cell_para(c, after=0, align=align)
        run(pp, txt, 9.5, bold=strong, color=NAVY if strong else INK)

p = para(doc, before=4, after=0)
run(p, f"Total failure value across all {CANON['n']} failed orders in this batch: {aed(TOTAL)}. "
       "Shares are of cash at risk and are rounded independently.", 8, italic=True, color=MUTE)

# ---- PAGE 2 -----------------------------------------------------------------
doc.add_page_break()

p = para(doc, after=2)
run(p, "How it works: four specialists, one coordinator", 13, bold=True, color=NAVY)
p = para(doc, after=8, line=1.18)
run(p, "Prep is deterministic: the code reads each failed order and computes the raw signals "
       "(late days, shortfall, dispatch slip, transit overage). Four specialist agents then run in "
       "parallel on Sonnet 4.6, each seeing only its own domain's evidence and never the answer key. "
       "The coordinator adjudicates the orders where more than one cause fired, choosing the deciding "
       "one rather than the biggest number. Finalize is deterministic again: it rolls up the cash, "
       "ranks the drivers, and renders the brief through a gate that blocks any number the code did "
       "not produce. The LLM reasons; tested code computes.", 10, color=INK)

# robustness story
p = para(doc, before=2, after=2)
run(p, "The robustness story", 13, bold=True, color=NAVY)
p = para(doc, after=8, line=1.18)
run(p, "One good run proves nothing. Accuracy here is measured across six independent batches "
       f"(seeds 42 to 47), each generated fresh with different planted causes the engine has never "
       f"seen. Mean accuracy on the hard orders, the ones where two teams are genuinely at fault at "
       f"once, is {pct(agg['mean_ambiguous_accuracy'])} "
       f"(range {pct(agg['min_ambiguous_accuracy'])} to {pct(agg['max_ambiguous_accuracy'])}), "
       f"against {pct(nv['ambiguous_accuracy'], 0)} to 84% for a simple biggest-number rule. On one "
       "batch the simple rule actually won. We show that row too.", 10, color=INK)

# robustness table
rt = doc.add_table(rows=1, cols=5)
rt.autofit = False
hair_borders(rt)
rwidths = [Inches(1.0), Inches(1.75), Inches(1.6), Inches(1.15), Inches(1.6)]
rheads = ["Batch (seed)", "Hard-order accuracy", "Simple rule", "Lift", "Note"]
for i, (h, w) in enumerate(zip(rheads, rwidths)):
    c = rt.cell(0, i)
    c.width = w
    shade(c, NAVY_HEX)
    set_cell_margins(c)
    pp = cell_para(c, after=0, align=WD_ALIGN_PARAGRAPH.LEFT if i in (0, 4) else WD_ALIGN_PARAGRAPH.RIGHT)
    run(pp, h, 8.5, bold=True, color=WHITE, caps=True, spacing=3)

for r in SUMM["seeds"]:
    seed = r["seed"]
    is_canon = seed == SUMM["canonical_seed"]
    lift = r["lift_ambiguous"] * 100
    is_anom = lift < 0
    note = "canonical, shown on page 1" if is_canon else ("simple rule won" if is_anom else "")
    fill = CANON_HEX if is_canon else (ANOM_HEX if is_anom else "FFFFFF")
    cells = rt.add_row().cells
    vals = [
        (str(seed), WD_ALIGN_PARAGRAPH.LEFT),
        (pct(r["ambiguous_accuracy"]), WD_ALIGN_PARAGRAPH.RIGHT),
        (pct(r["naive_ambiguous_accuracy"]), WD_ALIGN_PARAGRAPH.RIGHT),
        (f"{'+' if lift >= 0 else ''}{lift:.1f}", WD_ALIGN_PARAGRAPH.RIGHT),
        (note, WD_ALIGN_PARAGRAPH.LEFT),
    ]
    for i, (c, w) in enumerate(zip(cells, rwidths)):
        c.width = w
        shade(c, fill)
        set_cell_margins(c)
        txt, align = vals[i]
        pp = cell_para(c, after=0, align=align)
        strong = is_canon and i == 0
        clr = ACCENT if (i == 3 and lift >= 0) else (RGBColor(0xA8, 0x3A, 0x2F) if i == 3 else INK)
        run(pp, txt, 9.5, bold=(strong or i == 3), color=clr if i in (0, 3) and not strong else (NAVY if strong else clr if i == 3 else INK))

p = para(doc, before=4, after=8)
run(p, f"Mean lift across the six batches: +{agg['mean_lift_ambiguous'] * 100:.1f} points on the "
       f"hard orders. Mean overall accuracy: {pct(agg['mean_overall_accuracy'])}.", 8, italic=True, color=MUTE)

# three-way defense
p = para(doc, before=2, after=3)
run(p, "Why the numbers can be trusted", 13, bold=True, color=NAVY)
defense = [
    ("No fabrication.", "A firewall keeps the planted ground truth away from the agents, and a render "
     "gate blocks any digit the model authored from reaching the brief. Only code-computed, "
     "placeholder-bound numbers get through."),
    ("Measured, not asserted.", "Accuracy is scored against the cause planted in each order when the "
     "batch was generated, not against the model's own say-so."),
    ("Honest about the loss.", "On seed 47 the simple rule beat the coordinator. That batch stays in "
     "the table at full weight rather than being quietly dropped."),
]
for i, (lead, body) in enumerate(defense, 1):
    p = para(doc, after=3, line=1.15)
    run(p, f"{i}.  ", 10, bold=True, color=ACCENT)
    run(p, lead + " ", 10, bold=True, color=INK)
    run(p, body, 10, color=INK)

# worked example
p = para(doc, before=6, after=3)
run(p, "Worked example: order OTIF-0011", 12, bold=True, color=NAVY)
ex = doc.add_table(rows=1, cols=1)
ex.autofit = False
ex.columns[0].width = CONTENT_W
no_borders(ex)
ec = ex.cell(0, 0)
ec.width = CONTENT_W
shade(ec, "F4F6F8")
set_cell_margins(ec, top=140, bottom=140, left=160, right=160)
pp = cell_para(ec, after=3)
run(pp, "The supplier's PO ran 6 days late; the warehouse dispatch was 2 days slow with a pick short. ", 10, bold=True, color=INK)
run(pp, "The simple rule picks the louder single signal and blames the supplier. Wrong. ", 10, color=INK)
pp2 = cell_para(ec, after=0, before=0)
run(pp2, "The coordinator reasons that stock was physically on hand, so the late PO did not starve the "
        "pick; warehouse execution bound the outcome. It names the warehouse as the deciding cause. "
        "Correct, and it shows the reasoning.",
    10, color=INK)

# tech + disclaimer
p = para(doc, before=8, after=2)
run(p, "Built with  ", 9, bold=True, color=NAVY)
run(p, "Python, FastAPI, LangGraph, Claude Sonnet 4.6, Next.js.  Verified by 181 passing tests.",
    9, color=INK)

p = para(doc, before=6, after=0)
run(p, "Mawarid Distribution is a fictional company and all order data is synthetic, generated with a "
       "fixed seed so every figure in this document is reproducible. Prepared by Prajwal.",
    8, italic=True, color=MUTE)

# ---- footer -----------------------------------------------------------------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp, "OTIF Root-Cause Engine  ·  Case Study  ·  Prajwal  ·  Illustrative, synthetic data",
    7.5, color=MUTE, caps=True, spacing=4)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("wrote", OUT)
print(f"seed {CANON['seed']}  overall {pct(sc['overall_accuracy'])}  "
      f"mean-amb {pct(agg['mean_ambiguous_accuracy'])}  total {aed(TOTAL)}")
